import io
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_para_docx(texto_md: str, titulo: str = "") -> bytes:
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # Estilo padrão
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    if titulo:
        h = doc.add_heading(titulo.replace("-", " ").title(), level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    anterior_vazia = False
    for linha in texto_md.splitlines():
        if linha.strip() == "":
            if not anterior_vazia:
                doc.add_paragraph("")
            anterior_vazia = True
            continue
        anterior_vazia = False

        if linha.startswith("#### "):
            _adicionar_heading(doc, linha[5:], level=4)
        elif linha.startswith("### "):
            _adicionar_heading(doc, linha[4:], level=3)
        elif linha.startswith("## "):
            _adicionar_heading(doc, linha[3:], level=2, cor=True)
        elif linha.startswith("# "):
            _adicionar_heading(doc, linha[2:], level=1, cor=True)
        elif linha.startswith("- ") or linha.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _adicionar_inline(p, linha[2:])
        elif re.match(r"^\d+\. ", linha):
            p = doc.add_paragraph(style="List Number")
            _adicionar_inline(p, re.sub(r"^\d+\. ", "", linha))
        else:
            p = doc.add_paragraph()
            _adicionar_inline(p, linha)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _adicionar_heading(doc, texto: str, level: int, cor: bool = False):
    """Heading com marcação inline interpretada (sem asteriscos literais)."""
    h = doc.add_heading("", level=level)
    _adicionar_inline(h, texto)
    if cor:
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return h


def _adicionar_inline(paragraph, texto: str):
    # Suporte básico a **negrito** e *itálico*
    partes = re.split(r"(\*\*.*?\*\*|\*.*?\*)", texto)
    for parte in partes:
        if parte.startswith("**") and parte.endswith("**"):
            run = paragraph.add_run(parte[2:-2])
            run.bold = True
        elif parte.startswith("*") and parte.endswith("*"):
            run = paragraph.add_run(parte[1:-1])
            run.italic = True
        else:
            paragraph.add_run(parte)
